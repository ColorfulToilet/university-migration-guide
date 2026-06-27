#!/usr/bin/env python3
"""Convert a structured Markdown report into a simple Word .docx file.

The script prefers python-docx when available. If it is not installed, it falls
back to a minimal built-in OOXML writer so the Skill can still produce a real .docx.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re
import sys
from xml.sax.saxutils import escape
import zipfile


Block = tuple[str, object]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a Word report from Markdown.")
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", help="Output .docx file")
    return parser.parse_args()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped) <= {"|", "-", ":", " "}


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown(text: str) -> list[Block]:
    blocks: list[Block] = []
    table_buffer: list[str] = []
    number_index = 1

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [parse_table_row(line) for line in table_buffer if not is_table_separator(line)]
        if rows:
            blocks.append(("table", rows))
        table_buffer = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue

        flush_table()

        if not line.strip():
            number_index = 1
            continue
        if line.startswith("# "):
            blocks.append(("heading", (0, line[2:].strip())))
        elif line.startswith("## "):
            blocks.append(("heading", (1, line[3:].strip())))
        elif line.startswith("### "):
            blocks.append(("heading", (2, line[4:].strip())))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line).strip()
            blocks.append(("number", (number_index, text)))
            number_index += 1
        else:
            number_index = 1
            blocks.append(("paragraph", line.strip()))

    flush_table()
    return blocks


def write_with_python_docx(blocks: list[Block], output_path: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        return False

    document = Document()
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            document.add_heading(str(text), level=int(level))
        elif kind == "paragraph":
            document.add_paragraph(str(payload))
        elif kind == "bullet":
            document.add_paragraph(str(payload), style="List Bullet")
        elif kind == "number":
            _, text = payload
            document.add_paragraph(str(text), style="List Number")
        elif kind == "table":
            rows = payload
            width = max(len(row) for row in rows) if rows else 0
            table = document.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            for row in rows:
                cells = table.add_row().cells
                for index in range(width):
                    cells[index].text = row[index] if index < len(row) else ""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return True


def w_text(text: str, bold: bool = False, size: int | None = None) -> str:
    props = ""
    if bold or size:
        props_parts = []
        if bold:
            props_parts.append("<w:b/>")
        if size:
            props_parts.append(f'<w:sz w:val="{size}"/>')
        props = "<w:rPr>" + "".join(props_parts) + "</w:rPr>"
    return f"<w:r>{props}<w:t xml:space=\"preserve\">{escape(text)}</w:t></w:r>"


def w_para(text: str, bold: bool = False, size: int | None = None) -> str:
    return f"<w:p>{w_text(text, bold=bold, size=size)}</w:p>"


def w_table(rows: list[list[str]]) -> str:
    out = ["<w:tbl>"]
    out.append(
        "<w:tblPr><w:tblBorders>"
        "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "</w:tblBorders></w:tblPr>"
    )
    for row in rows:
        out.append("<w:tr>")
        for cell in row:
            out.append(f"<w:tc><w:p>{w_text(cell)}</w:p></w:tc>")
        out.append("</w:tr>")
    out.append("</w:tbl>")
    return "".join(out)


def write_minimal_docx(blocks: list[Block], output_path: Path) -> None:
    body_parts: list[str] = []
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            size = 36 if int(level) == 0 else 30 if int(level) == 1 else 26
            body_parts.append(w_para(str(text), bold=True, size=size))
        elif kind == "paragraph":
            body_parts.append(w_para(str(payload)))
        elif kind == "bullet":
            body_parts.append(w_para("• " + str(payload)))
        elif kind == "number":
            number, text = payload
            body_parts.append(w_para(f"{number}. {text}"))
        elif kind == "table":
            body_parts.append(w_table(payload))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(body_parts)
        + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        "</w:body></w:document>"
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.is_file():
        print(f"找不到输入文件：{input_path}", file=sys.stderr)
        return 1

    blocks = parse_markdown(input_path.read_text(encoding="utf-8"))
    if not write_with_python_docx(blocks, output_path):
        write_minimal_docx(blocks, output_path)
    print(f"已生成：{output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
